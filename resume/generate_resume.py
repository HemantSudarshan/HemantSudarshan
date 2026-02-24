import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Style
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def format_section_heading(doc, text):
    heading = doc.add_paragraph()
    run = heading.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    heading.paragraph_format.space_after = Pt(4)
    heading.paragraph_format.space_before = Pt(10)
    
    pPr = heading._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_bullet(doc, text, bold_text=None, italic_text=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    if bold_text:
        run = p.add_run(bold_text)
        run.font.bold = True
        run.font.size = Pt(10)
    if italic_text:
        run = p.add_run(italic_text)
        run.font.italic = True
        run.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def create_resume():
    doc = Document()
    
    # 0.5 inch margins all around for ATS styling and max content fit
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # HEADER (ATS Optimized text fields without emojis/tables)
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(4)
    name_run = header.add_run('HEMANT SUDARSHAN\n')
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    
    title_run = header.add_run('AI Localization Consultant | Open Source Contributor | Production AI Systems Developer\n')
    title_run.font.size = Pt(11)
    
    contact_run = header.add_run('Bengaluru, India | collabwithhemantgenai@gmail.com | ')
    contact_run.font.size = Pt(10)
    
    add_hyperlink(header, 'LinkedIn', 'https://linkedin.com/in/hemant-sudarshan-01633928a')
    header.add_run(' | ').font.size = Pt(10)
    add_hyperlink(header, 'GitHub & Portfolio', 'https://github.com/HemantSudarshan')
    header.add_run(' | ').font.size = Pt(10)
    add_hyperlink(header, 'HuggingFace', 'https://huggingface.co/Hemantxai')
    
    # SUMMARY
    format_section_heading(doc, 'PROFESSIONAL SUMMARY')
    summary = doc.add_paragraph(
        'Driven and innovative Computer Science graduate with extensive hands-on experience building and deploying 6+ full-cycle production-ready AI systems. '
        'Highly skilled in designing robust retrieval-augmented generation (RAG) pipelines, complex multi-agent AI architectures, and managing end-to-end cloud deployments using Docker, FastAPI, and CI/CD pipelines. '
        'Demonstrated success in optimizing large language models (LLMs) to mitigate AI hallucinations, managing scalable vector search systems (Weaviate, Qdrant), and engineering robust enterprise AI applications. '
        'Offers a strong foundation in open-source engineering with impactful contributions to performance tooling, holds 1 technology patent, and is distinguished by a proactive approach to solving real-world business challenges through scalable AI solutions.'
    )
    summary.paragraph_format.space_after = Pt(6)
    summary.paragraph_format.space_before = Pt(4)
    for run in summary.runs:
        run.font.size = Pt(10)
        
    # TECHNICAL SKILLS
    format_section_heading(doc, 'TECHNICAL SKILLS')
    add_bullet(doc, ': LangChain, LangGraph, PyTorch, TensorFlow, RAG Pipelines', bold_text='AI/ML Frameworks')
    add_bullet(doc, ': GPT-4o, Claude, Gemini, Llama, Groq, SambaNova', bold_text='LLMs & APIs')
    add_bullet(doc, ': Python, FastAPI, Flask, Docker, MongoDB, Redis', bold_text='Backend & DevOps')
    add_bullet(doc, ': Weaviate, Qdrant, Pinecone, ChromaDB', bold_text='Databases & Vector DBs')
    add_bullet(doc, ': Git, GitHub Actions, CI/CD, AWS, Railway', bold_text='DevOps')

    # EXPERIENCE
    format_section_heading(doc, 'PROFESSIONAL EXPERIENCE')
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('AI Operations & Localization Consultant')
    r1.font.bold = True
    r1.font.size = Pt(11)
    p.add_run(' | Pratilipi Comics | Oct 2024 – Present').font.size = Pt(10.5)
    add_bullet(doc, 'Leading AI-driven localization initiatives for Indic language content (Hindi, Telugu, Kannada)')
    add_bullet(doc, 'Building GenAI pipelines for multi-language content market (600M+ potential users)')
    add_bullet(doc, 'Optimizing content delivery workflows with custom AI tools (40% efficiency improvement)')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('Localization Operations & Gen AI Intern')
    r1.font.bold = True
    r1.font.size = Pt(11)
    p.add_run(' | Pratilipi Comics | Feb 2024 – Oct 2024').font.size = Pt(10.5)
    add_bullet(doc, 'Developed GenAI tools for multi-language content localization (3+ languages)')
    add_bullet(doc, 'Automated translation and localization workflows using LLM-based pipelines')
    add_bullet(doc, 'Reduced manual localization effort by 40% through AI workflow optimization')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('Machine Learning Intern')
    r1.font.bold = True
    r1.font.size = Pt(11)
    p.add_run(' | AppiVa Software Pvt Ltd (NASSCOM 10K Startup) | Aug 2023 – Nov 2023').font.size = Pt(10.5)
    add_bullet(doc, 'Built data preprocessing and analytics pipelines')
    add_bullet(doc, 'Developed and optimized ML models for production')
    add_bullet(doc, 'First hands-on experience with production ML systems')

    # PROJECTS
    format_section_heading(doc, 'KEY PROJECTS')
    
    # Proj 1
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('Agentic Inventory Restocking Service | ')
    r1.font.bold = True
    r1.font.size = Pt(11)
    p.add_run('LangGraph, Gemini AI, FastAPI, MongoDB | ').font.size = Pt(10)
    add_hyperlink(p, 'GitHub', 'https://github.com/HemantSudarshan/Agentic-Inventory-Restocking-Service')
    add_bullet(doc, 'Engineered an advanced multi-agent workflow to autonomously analyze product demand patterns via time-series forecasting, mimicking human decision-making for complex inventory routing.')
    add_bullet(doc, 'Created a highly resilient dual-model fallback system (Gemini 2.0 primary + Groq backup), achieving 99.9% availability, and driving a 95% reduction in false restock alarms and systemic over-purchasing.')
    add_bullet(doc, 'Developed an interactive, high-performance administrative dashboard utilizing FastAPI, Redis cache layers for rapid session lookup, and fully typed request validation via Pydantic to ensure reliable data operations.')

    # Proj 2
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('Compliance-GPT: Enterprise RAG System | ')
    r1.font.bold = True
    r1.font.size = Pt(11)
    p.add_run('Weaviate, Groq, FastAPI, Docker | ').font.size = Pt(10)
    add_hyperlink(p, 'GitHub', 'https://github.com/HemantSudarshan/Compliance-GPT')
    add_bullet(doc, 'Modeled and deployed a zero-hallucination compliance assistant that significantly reduced research times, yielding comprehensive, legally citation-backed sub-2s responses for GDPR, HIPAA, and CCPA queries.')
    add_bullet(doc, 'Employed hybrid vector search infrastructure via Weaviate (merging BM25 keyword matching with Semantic Vectors), enforcing strict chunk-level citation extraction to guarantee the exact origin of legal facts.')
    add_bullet(doc, 'Containerized the complete search application using Docker, enforcing enterprise-level API limits with rate throttling, rigid CORS configuration blocks, and comprehensive HTTP security headers.')

    # Proj 3
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('AudioRAG Enterprise: Multi-Tenant Audio Analytics | ')
    r1.font.bold = True
    r1.font.size = Pt(11)
    p.add_run('AssemblyAI, Qdrant, SambaNova, PostgreSQL | ').font.size = Pt(10)
    add_hyperlink(p, 'GitHub', 'https://github.com/HemantSudarshan/AudioRaG_FTSambanova')
    add_bullet(doc, 'Constructed a robust audio-ingestion pipeline featuring real-time, highly accurate speaker diarization and automated HIPAA/GDPR text redaction using AssemblyAI integrations.')
    add_bullet(doc, 'Optimized query latency and retrieval accuracy using domain-aware BGE-Large embeddings on Qdrant vector databases, supplemented effectively by scalable PostgreSQL audit logs.')

    # Proj 4
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('TruthTracker (AntiAi Deepfake Platform) | ')
    r1.font.bold = True
    r1.font.size = Pt(11)
    p.add_run('PyTorch, EfficientNet-B4, Gradio | ').font.size = Pt(10)
    add_hyperlink(p, 'GitHub', 'https://github.com/HemantSudarshan/AntiAi')
    add_bullet(doc, 'Designed a sophisticated computer vision application utilizing an EfficientNet-B4 network architecture to systematically detect and flag audio-visual deepfakes alongside known fake news content signatures.')
    add_bullet(doc, 'Wrapped the inference logic within an interactive Gradio interface, enabling seamless real-time user testing and demonstration of the media verification model.')

    # Proj 5
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('AI Real Estate Agent | ')
    r1.font.bold = True
    r1.font.size = Pt(11)
    p.add_run('Gemini AI, Firecrawl, Redis | ').font.size = Pt(10)
    add_hyperlink(p, 'GitHub', 'https://github.com/HemantSudarshan/-AI-Powered-Real-Estate-Agent-Automating-Property-Search-Investment-Insights-')
    add_bullet(doc, 'Implemented an autonomous property search assistant utilizing a multi-agent framework orchestrating Google Gemini endpoints, heavily integrated with Firecrawl logic for targeted web scraping.')
    add_bullet(doc, 'Streamlined the real estate search process by aggregating and ranking relevant property descriptions based on intelligent semantic analysis of dynamic user preferences.')

    # OPEN-SOURCE CONTRIBUTIONS
    format_section_heading(doc, 'OPEN-SOURCE CONTRIBUTIONS')
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('openclaw.ai (PR #37 - Merged) | JavaScript Performance Optimization')
    r1.font.bold = True
    r1.font.size = Pt(11)
    add_bullet(doc, 'Overhauled the core DOM manipulation scripts by aggressively caching variables and eliminating repeated query lookups, integrating extensive null safety to cut execution overhead substantially.')
    add_bullet(doc, 'Standardized the repository’s dependency management to prevent versioning conflicts, significantly bolstering the application’s crash resilience under heavy load.')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('Kreuzberg (PR #389 - Merged) | Document Extraction Python API')
    r1.font.bold = True
    r1.font.size = Pt(11)
    add_bullet(doc, 'Discovered and definitively remedied synchronization flaws within the core document extraction migration guides, directly clarifying async interfaces for the wider developer community.')
    add_bullet(doc, 'Upgraded deprecated example scripts to executable Python output flows, reinforcing documentation accuracy—a critical component for external API trust, reliability, and library adoption.')

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('docling-project/docling (PR #3022 - Merged) | DOCX Backend Bug Fix (IBM Open Source)')
    r1.font.bold = True
    r1.font.size = Pt(11)
    add_bullet(doc, 'Identified and fixed a crash in the DOCX parsing backend where Path(c.address) raised a TypeError when c.address is None for internal bookmark hyperlinks (e.g., TOC entries, cross-references), causing complete document conversion failure.')
    add_bullet(doc, 'Introduced a one-line defensive guard (hyperlink = Path(c.address) if c.address else None) following the existing null-safety pattern on the adjacent line; all 12 existing DOCX backend tests continue to pass unchanged.')
    add_bullet(doc, 'Added regression test test_hyperlink_with_none_address that programmatically constructs a DOCX via raw XML manipulation to reproduce the exact failure case, contributing 57 lines across msword_backend.py and test_backend_msword.py.')

    # RESEARCH, ACHIEVEMENTS & EDUCATION
    format_section_heading(doc, 'RESEARCH & ACHIEVEMENTS')
    add_bullet(doc, 'A System for Providing Security Using a Plurality of Factors for IoT Gadgets (Indian Patent No. 202341040746)', bold_text='Patent Filed: ')
    add_bullet(doc, 'Discovering Insights into Heart Health: A Survey of Data Mining and Machine Learning Methods (Presented at ICCICCT-2023 NICHE)', bold_text='Conference Publication: ')
    add_bullet(doc, 'Survey of AI-Driven Platforms for Welfare and Emergency Services: Gaps, Architectures and the Case for Unified Systems — GRENZE International Journal of Engineering and Technology (GIJET), Vol. 11, Issue 2, Pages 9911–9916, 2025. Co-authors: Lavanya Ramkumar, Afsha R, Vinayaka VM.', bold_text='Journal Publication: ')
    
    # Adding separate Education section to match the structure listed in README.md
    format_section_heading(doc, 'EDUCATION')
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run('B.Tech in Computer Science & Technology')
    r1.font.bold = True
    r1.font.size = Pt(11)
    p.add_run(' | Dayananda Sagar University, Bangalore | 2021 – 2025 (First Class)').font.size = Pt(10.5)
    
    # Save directly to the target location
    output_path = r'c:\Python Project\HemantSudarshan\resume\Hemant_Sudarshan_AI_ML_Engineer_Resume.docx'
    
    # Fallback to create directory if not exists, though it should explicitly be there
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"✅ Detailed ATS-Friendly Resume generated successfully: {output_path}")

if __name__ == "__main__":
    create_resume()
